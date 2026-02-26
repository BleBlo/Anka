from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set margins - standard 0.5" all around
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_spacing(paragraph, before=0, after=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)

# ============================================================================
# HEADER
# ============================================================================
name = doc.add_paragraph()
name_run = name.add_run("SAIF KHALFAN ALMAZROUEI")
name_run.bold = True
name_run.font.size = Pt(18)
name_run.font.name = 'Calibri'
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(name, before=0, after=3)

contact = doc.add_paragraph()
contact_run = contact.add_run("Madison, WI | skalmazrouei@wisc.edu | +1 (312) 366-6300 | +971 50 199 9047")
contact_run.font.size = Pt(10)
contact_run.font.name = 'Calibri'
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(contact, before=0, after=8)
add_horizontal_line(contact)

# ============================================================================
# EDUCATION
# ============================================================================
edu_header = doc.add_paragraph()
edu_run = edu_header.add_run("EDUCATION")
edu_run.bold = True
edu_run.font.size = Pt(12)
set_spacing(edu_header, before=8, after=4)
add_horizontal_line(edu_header)

edu = doc.add_paragraph()
edu_bold = edu.add_run("University of Wisconsin-Madison")
edu_bold.bold = True
edu_bold.font.size = Pt(11)
edu.add_run(" | B.S. Data Science, Computer Science, Information Science").font.size = Pt(11)
set_spacing(edu, before=4, after=2)

edu2 = doc.add_paragraph()
edu2.add_run("Expected May 2026 | GPA: 3.6/4.0 | Minor: Digital Studies").font.size = Pt(10)
set_spacing(edu2, before=0, after=2)

edu3 = doc.add_paragraph()
edu3.add_run("Relevant Coursework: Artificial Intelligence (A), Linear Algebra (A), Algorithms, Machine Learning, Human-Computer Interaction, Data Science Programming, Machine Organization").font.size = Pt(10)
set_spacing(edu3, before=0, after=6)

# ============================================================================
# RESEARCH
# ============================================================================
research_header = doc.add_paragraph()
research_run = research_header.add_run("RESEARCH")
research_run.bold = True
research_run.font.size = Pt(12)
set_spacing(research_header, before=6, after=4)
add_horizontal_line(research_header)

research_title = doc.add_paragraph()
rt_bold = research_title.add_run("Anka: A Domain-Specific Language for Reliable LLM Code Generation")
rt_bold.bold = True
rt_bold.font.size = Pt(11)
research_title.add_run(" | arXiv 2024").font.size = Pt(11)
set_spacing(research_title, before=4, after=2)

bullets = [
    "Designed and implemented a novel domain-specific language that achieves a 40 percentage point accuracy improvement over Python on multi-step code generation tasks, addressing systematic LLM failures in complex pipeline construction",
    "Built complete language toolchain from scratch: Lark-based parser with 99 grammar rules, semantic analyzer with type checking, and tree-walking interpreter, validated by 322 automated tests achieving comprehensive coverage",
    "Demonstrated zero-shot DSL learning: Claude 3.5 Haiku achieves 95.8% task accuracy and 99.9% parse success on a novel language with zero prior training exposure, learning entirely from in-context examples",
    "Validated findings across model families (GPT-4o-mini: +26.7pp on complex pipelines); open-sourced complete implementation, 100-task benchmark suite, and evaluation framework for reproducibility"
]

for bullet in bullets:
    bp = doc.add_paragraph()
    bp.add_run("• " + bullet).font.size = Pt(10)
    bp.paragraph_format.left_indent = Inches(0.2)
    set_spacing(bp, before=2, after=2)

# ============================================================================
# PATENTS
# ============================================================================
patent_header = doc.add_paragraph()
patent_run = patent_header.add_run("PATENTS & INTELLECTUAL PROPERTY")
patent_run.bold = True
patent_run.font.size = Pt(12)
set_spacing(patent_header, before=8, after=4)
add_horizontal_line(patent_header)

patent_title = doc.add_paragraph()
pt_bold = patent_title.add_run("USPTO Provisional Patent Applications")
pt_bold.bold = True
pt_bold.font.size = Pt(11)
patent_title.add_run(" | November 2025").font.size = Pt(11)
set_spacing(patent_title, before=4, after=2)

patent_bullets = [
    "Multi-Agent LLM Trading Architecture: Autonomous trading strategy generation system using structured debate protocols between specialized AI agents, enabling self-improving investment decisions through adversarial refinement",
    "Cost-Optimized Single-Model Trading System: Novel persona-switching approach achieving comparable performance to multi-agent systems at 70-80% lower computational cost, democratizing AI-powered trading capabilities",
    "Adaptive Multi-Dimensional Data Influence Scoring: Dynamic optimization framework for AI-powered decision-making that continuously recalibrates feature importance across diverse application domains"
]

for bullet in patent_bullets:
    bp = doc.add_paragraph()
    bp.add_run("• " + bullet).font.size = Pt(10)
    bp.paragraph_format.left_indent = Inches(0.2)
    set_spacing(bp, before=2, after=2)

# ============================================================================
# EXPERIENCE
# ============================================================================
exp_header = doc.add_paragraph()
exp_run = exp_header.add_run("EXPERIENCE")
exp_run.bold = True
exp_run.font.size = Pt(12)
set_spacing(exp_header, before=8, after=4)
add_horizontal_line(exp_header)

# ADIA 2024
exp1_title = doc.add_paragraph()
e1_bold = exp1_title.add_run("Abu Dhabi Investment Authority")
e1_bold.bold = True
e1_bold.font.size = Pt(11)
exp1_title.add_run(" | Central Investment Services Intern | Summer 2024").font.size = Pt(11)
set_spacing(exp1_title, before=4, after=2)

exp1_bullets = [
    "Spearheaded development of PRISM (PDF Retrieval, Interaction & Smart Management): AI-powered tool that revolutionized investment data processing, reducing extraction time from 25.5 hours to 2 minutes per document (96% reduction) while achieving 98%+ accuracy",
    "Engineered comprehensive end-to-end pipeline integrating Python, Streamlit, PyMuPDF, and LLM APIs to automate complex data analysis from 200+ sources; successfully deployed across 6 departments with full documentation for ADIA-wide integration"
]

for bullet in exp1_bullets:
    bp = doc.add_paragraph()
    bp.add_run("• " + bullet).font.size = Pt(10)
    bp.paragraph_format.left_indent = Inches(0.2)
    set_spacing(bp, before=2, after=2)

# ADIA 2023
exp2_title = doc.add_paragraph()
e2_bold = exp2_title.add_run("Abu Dhabi Investment Authority")
e2_bold.bold = True
e2_bold.font.size = Pt(11)
exp2_title.add_run(" | IT Research Intern | Summer 2023").font.size = Pt(11)
set_spacing(exp2_title, before=6, after=2)

exp2_bullets = [
    "Collaborated with Goldman Sachs team to enhance financial analysis capabilities by developing an interactive dashboard for visualizing and analyzing investment portfolios, significantly improving data accessibility for decision-makers",
    "Contributed to optimization of investment strategies through detailed basket analysis, providing actionable insights that informed portfolio management decisions across multiple asset classes"
]

for bullet in exp2_bullets:
    bp = doc.add_paragraph()
    bp.add_run("• " + bullet).font.size = Pt(10)
    bp.paragraph_format.left_indent = Inches(0.2)
    set_spacing(bp, before=2, after=2)

# ============================================================================
# LEADERSHIP
# ============================================================================
lead_header = doc.add_paragraph()
lead_run = lead_header.add_run("LEADERSHIP & PROGRAMS")
lead_run.bold = True
lead_run.font.size = Pt(12)
set_spacing(lead_header, before=8, after=4)
add_horizontal_line(lead_header)

lead_title = doc.add_paragraph()
l_bold = lead_title.add_run("Masdar Youth 4 Sustainability")
l_bold.bold = True
l_bold.font.size = Pt(11)
lead_title.add_run(" | Future Sustainability Leader & Ambassador | 2019-2024").font.size = Pt(11)
set_spacing(lead_title, before=4, after=2)

lead_bullets = [
    "Selected for elite Sustainability Ambassadors (2019-2020) and Future Sustainability Leaders (2023-2024) cohorts, with only 50 participants chosen annually from global applicant pool",
    "Led workshops at Abu Dhabi Sustainability Week; collaborated with international policymakers on SDG implementation and clean energy transition strategies within Masdar's 41,000+ member Y4S network"
]

for bullet in lead_bullets:
    bp = doc.add_paragraph()
    bp.add_run("• " + bullet).font.size = Pt(10)
    bp.paragraph_format.left_indent = Inches(0.2)
    set_spacing(bp, before=2, after=2)

# ADIA Early Prep
lead2_title = doc.add_paragraph()
l2_bold = lead2_title.add_run("Abu Dhabi Investment Authority Early Preparation Program")
l2_bold.bold = True
l2_bold.font.size = Pt(11)
lead2_title.add_run(" | 2017-2020").font.size = Pt(11)
set_spacing(lead2_title, before=6, after=2)

lead2_bp = doc.add_paragraph()
lead2_bp.add_run("• Participated in career development seminars and international summer programs at University of Pennsylvania and in China").font.size = Pt(10)
lead2_bp.paragraph_format.left_indent = Inches(0.2)
set_spacing(lead2_bp, before=2, after=2)

# ============================================================================
# SKILLS
# ============================================================================
skills_header = doc.add_paragraph()
skills_run = skills_header.add_run("TECHNICAL SKILLS")
skills_run.bold = True
skills_run.font.size = Pt(12)
set_spacing(skills_header, before=8, after=4)
add_horizontal_line(skills_header)

skills = doc.add_paragraph()
s1_bold = skills.add_run("Languages & ML: ")
s1_bold.bold = True
s1_bold.font.size = Pt(10)
skills.add_run("Python, PyTorch, TensorFlow, scikit-learn, Hugging Face Transformers, LangChain, OpenAI API, Anthropic API").font.size = Pt(10)
set_spacing(skills, before=4, after=2)

skills2 = doc.add_paragraph()
s2_bold = skills2.add_run("Infrastructure & Tools: ")
s2_bold.bold = True
s2_bold.font.size = Pt(10)
skills2.add_run("NumPy, Pandas, Git, Docker, AWS, CUDA, Linux, Streamlit, Jupyter, SQL").font.size = Pt(10)
set_spacing(skills2, before=2, after=2)

skills3 = doc.add_paragraph()
s3_bold = skills3.add_run("Compiler & Language Design: ")
s3_bold.bold = True
s3_bold.font.size = Pt(10)
skills3.add_run("Lark (EBNF parsing), Abstract Syntax Trees, Semantic Analysis, Type Systems, Tree-Walking Interpreters").font.size = Pt(10)
set_spacing(skills3, before=2, after=0)

# Save
doc.save('Almazrouei_Saif_Resume_New.docx')
print("Resume created successfully: Almazrouei_Saif_Resume_New.docx")
